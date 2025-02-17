# !/usr/bin/env python
# coding= utf-8
"""
@Project         : e2w
@File            : e2w.py
@Author          : yum <richardminyu@foxmail.com>
@Date            : 2025/02/17 19:41
@Description     : transform excel to word
"""

import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import datetime


def get_excel_content():
    wb_obj = openpyxl.load_workbook("./fj4.xlsx")
    sheet_obj = wb_obj.active
    max_row = sheet_obj.max_row
    res = []
    c = 0

    for i in range(2, max_row + 1):
        c += 1
        # print(c, res)
        res.append({"照片号": c,
                    "物品名称": sheet_obj.cell(row=i, column=2).value,
                    "拍摄地点": sheet_obj.cell(row=i, column=3).value,
                    "对应序号": sheet_obj.cell(row=i, column=1).value,
                    "拍摄时间": ""})
    return res


def set_word_content():
    document = Document()
    document.styles['Normal'].font.size = Pt(14)
    # 英文字体
    document.styles['Normal'].font.name = u"Times New Roman"
    # 中文字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.add_heading("附件7：", level=0)
    document.add_heading("安全生产费用申报附表（图）说明", level=1)
    result = get_excel_content()
    c = 0
    try:
        for i in result:
            c += 1
            p = document.add_paragraph()
            print("000", i, c)
            # print('000', i["照片号"])
            # # print('000', i["物品名称"])
            # # print('000', i["拍摄地点"])
            if i["照片号"] and i["物品名称"] and i["拍摄地点"] and i["对应序号"]:
                p.add_run("照片号：" + str(i["照片号"]) + "\n")
                p.add_run("物品名称：" + str(i["物品名称"]) + "\n")
                p.add_run("拍摄地点：" + str(i["拍摄地点"]) + "\n")
                p.add_run("对应序号：" + str(i["对应序号"]) + "\n")
                p.add_run("拍摄时间：" + (i["拍摄时间"] or "") + "\n")
            if c % 2 == 0:
                document.add_page_break()
            # print(i)
        nt = datetime.datetime.now().strftime('%Y-%m-%d-%H-%M-%S')
        document.save(f"result_{nt}.docx")
    except Exception as e:
        print("error: ", e)


if __name__ == "__main__":
    set_word_content()
